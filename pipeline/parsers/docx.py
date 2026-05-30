"""
DOCX parser using python-docx.
Extracts paragraph text. Tables are flattened cell-by-cell.
No page number metadata available from docx format -- page is always 1.
"""

from pathlib import Path


def extract(filepath: str) -> list[dict]:
    """
    Returns [{"page": 1, "text": str}] -- single entry, full document text.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(filepath)

    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip()
                                  for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    full_text = "\n".join(parts)
    if not full_text.strip():
        return []

    return [{"page": 1, "text": full_text}]
